
"""
Word 文档操作模块：
提供 docx/doc 文档的创建、读取、更新、删除，并支持基础格式设置。
"""
import os
import sys

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
if project_root not in sys.path:
    sys.path.append(project_root)

from ai_files_tools.ai_files_read import resolve_desktop_path, resolve_target_path, validate_path_security


def create_docx_file(file_name, content=None, parent_path=None, style=None):
    """
    创建 Word 文档，支持写入内容与基础格式设置。

    Args:
        file_name (str): 文件名，支持 .docx 或 .doc。
        content (str|list|dict, optional): 文本内容或结构化内容。
        parent_path (str, optional): 父目录路径，不传默认桌面。
        style (dict, optional): 字体样式，支持 font_name/font_size/color。

    Returns:
        dict: 操作结果信息。
    """
    if not file_name:
        return {"success": False, "reason": "file_name_empty", "message": "文件名不能为空"}
    target_parent = _resolve_parent_path(parent_path)
    if not target_parent:
        return {"success": False, "reason": "parent_path_invalid", "message": "无法定位父目录"}
    safe_name = _ensure_extension(str(file_name).strip(), ".docx")
    full_path = os.path.join(target_parent, safe_name)
    
    if os.path.exists(full_path):
        return {"success": False, "reason": "already_exists", "path": full_path, "message": "目标文件已存在"}
    try:
        extension = _resolve_extension(safe_name)
        if extension == ".doc":
            result = _create_doc_file_with_word(full_path, content, style)
            if not result.get("success"):
                return result
        else:
            doc = _create_document()
            _append_content(doc, content, style)
            doc.save(full_path)
        return {"success": True, "path": full_path, "message": "创建成功"}
    except Exception as e:
        return {"success": False, "reason": "create_failed", "error": str(e), "message": f"创建失败: {str(e)}"}


def read_docx_file(file_path):
    """
    读取 Word 文档的纯文本内容。

    Args:
        file_path (str): docx/doc 文件路径或名称。

    Returns:
        dict: 读取结果信息。
    """
    normalized = _normalize_existing_path(file_path, (".docx", ".doc"))
    if not normalized:
        return {"success": False, "reason": "path_invalid_or_not_found", "message": "路径无效或文件不存在"}
    try:
        extension = _resolve_extension(normalized)
        if extension == ".doc":
            text = _read_doc_with_word(normalized)
        else:
            doc = _load_document(normalized)
            text = "\n".join([p.text for p in doc.paragraphs])
        return {"success": True, "path": normalized, "content": text}
    except Exception as e:
        return {"success": False, "reason": "read_failed", "error": str(e), "message": f"读取失败: {str(e)}"}


def update_docx_content(file_path, content, mode="replace", style=None):
    """
    更新 Word 文档内容，支持替换、追加、删除，并支持基础格式设置。

    Args:
        file_path (str): docx/doc 文件路径或名称。
        content (str|list|dict): 文本内容或结构化内容。
        mode (str): replace/append/remove。
        style (dict, optional): 字体样式，支持 font_name/font_size/color。

    Returns:
        dict: 更新结果信息。
    """
    normalized = _normalize_existing_path(file_path, (".docx", ".doc"))
    if not normalized:
        return {"success": False, "reason": "path_invalid_or_not_found", "message": "路径无效或文件不存在"}
    mode = str(mode or "replace").lower()
    try:
        extension = _resolve_extension(normalized)
        if extension == ".doc":
            return _update_doc_with_word(normalized, content, mode, style)
        doc = _load_document(normalized)
        if mode == "replace":
            _clear_document(doc)
            _append_content(doc, content, style)
        elif mode == "append":
            if content is None:
                return {"success": False, "reason": "content_empty", "message": "追加内容不能为空"}
            _append_content(doc, content, style)
        elif mode == "remove":
            if not content:
                return {"success": False, "reason": "content_empty", "message": "要删除的内容不能为空"}
            removed = _remove_paragraphs(doc, str(content))
            if removed == 0:
                return {"success": False, "reason": "content_not_found", "message": "未找到要删除的内容"}
        else:
            return {"success": False, "reason": "mode_invalid", "message": "不支持的更新模式"}
        doc.save(normalized)
        return {"success": True, "path": normalized, "message": "内容已更新", "mode": mode}
    except Exception as e:
        return {"success": False, "reason": "update_failed", "error": str(e), "message": f"更新失败: {str(e)}"}


def delete_docx_file(file_path):
    """
    删除 Word 文档。

    Args:
        file_path (str): docx/doc 文件路径或名称。

    Returns:
        dict: 删除结果信息。
    """
    normalized = _normalize_existing_path(file_path, (".docx", ".doc"))
    if not normalized:
        return {"success": False, "reason": "path_invalid_or_not_found", "message": "路径无效或文件不存在"}
    try:
        os.remove(normalized)
        return {"success": True, "path": normalized, "message": "删除成功"}
    except Exception as e:
        return {"success": False, "reason": "delete_failed", "error": str(e), "message": f"删除失败: {str(e)}"}


def _resolve_parent_path(parent_path):
    if parent_path:
        target_parent = resolve_target_path(parent_path)
    else:
        target_parent = resolve_desktop_path()
    
    if not target_parent or not os.path.exists(target_parent) or not os.path.isdir(target_parent):
        return None
    return os.path.abspath(os.path.normpath(target_parent))


def _ensure_extension(name, extension):
    lower_name = name.lower()
    if lower_name.endswith(".docx") or lower_name.endswith(".doc"):
        return name
    return f"{name}{extension}"


def _normalize_existing_path(path, extensions):
    if not path:
        return None
    if isinstance(extensions, str):
        extensions = (extensions,)
    
    # 使用统一的路径解析
    resolved = resolve_target_path(path)
    if not resolved or not os.path.exists(resolved):
        return None
        
    resolved = os.path.abspath(os.path.normpath(resolved))
    
    if not any(resolved.lower().endswith(ext) for ext in extensions):
        return None
    return resolved


def _create_document():
    from docx import Document
    return Document()


def _load_document(path):
    from docx import Document
    return Document(path)


def _append_content(doc, content, style=None):
    """
    按内容与样式追加到 docx 文档。
    """
    if content is None:
        return
    blocks = _normalize_content_blocks(content)
    if not blocks:
        _append_paragraph_with_style(doc, "", style)
        return
    for block in blocks:
        text = block.get("text", "")
        merged_style = _merge_styles(style, block.get("style"))
        lines = str(text).splitlines()
        if not lines:
            _append_paragraph_with_style(doc, "", merged_style)
            continue
        for line in lines:
            _append_paragraph_with_style(doc, line, merged_style)


def _append_paragraph_with_style(doc, text, style):
    """
    追加单段落文本并应用样式。
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _apply_run_style(run, style)


def _normalize_content_blocks(content):
    """
    统一内容结构为可处理的块列表。
    """
    if isinstance(content, (list, tuple)):
        blocks = []
        for item in content:
            blocks.extend(_normalize_single_block(item))
        return blocks
    return _normalize_single_block(content)


def _normalize_single_block(item):
    """
    标准化单个内容块。
    """
    if isinstance(item, dict):
        text = item.get("text")
        if text is None and "content" in item:
            text = item.get("content")
        if text is None and "lines" in item:
            lines = item.get("lines")
            if isinstance(lines, (list, tuple)):
                text = "\n".join([str(line) for line in lines])
        return [{"text": "" if text is None else str(text), "style": item.get("style")}]
    return [{"text": "" if item is None else str(item), "style": None}]


def _content_to_plain_text(content):
    """
    将结构化内容转换为纯文本，用于 .doc 写入。
    """
    blocks = _normalize_content_blocks(content)
    if not blocks:
        return ""
    lines = []
    for block in blocks:
        text = block.get("text", "")
        text_lines = str(text).splitlines()
        if not text_lines:
            lines.append("")
            continue
        lines.extend(text_lines)
    return "\n".join(lines)


def _merge_styles(base_style, extra_style):
    """
    合并样式配置，后者优先级更高。
    """
    merged = _normalize_style(base_style)
    extra = _normalize_style(extra_style)
    for key, value in extra.items():
        if value is not None:
            merged[key] = value
    return merged


def _normalize_style(style):
    """
    将输入样式统一为标准字段。
    """
    if not isinstance(style, dict):
        return {}
    return {
        "font_name": style.get("font_name") or style.get("font") or style.get("name"),
        "font_size": style.get("font_size") or style.get("size"),
        "color": style.get("color") or style.get("font_color")
    }


def _apply_run_style(run, style):
    """
    将样式应用到 docx 的文字 run。
    """
    style = _normalize_style(style)
    if not style:
        return
    from docx.shared import Pt, RGBColor

    font_name = style.get("font_name")
    if font_name:
        run.font.name = str(font_name)

    font_size = _parse_font_size(style.get("font_size"))
    if font_size:
        run.font.size = Pt(font_size)

    color = _parse_color(style.get("color"))
    if color:
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])


def _parse_font_size(value):
    """
    解析字号为浮点数。
    """
    if value is None:
        return None
    try:
        size = float(value)
        return size if size > 0 else None
    except Exception:
        return None


def _parse_color(value):
    """
    解析颜色输入，返回 (r, g, b) 元组。
    """
    if value is None:
        return None
    if isinstance(value, (list, tuple)) and len(value) == 3:
        return _normalize_rgb_tuple(value)
    if isinstance(value, str):
        raw = value.strip()
        lower = raw.lower()
        if lower.startswith("rgb(") and lower.endswith(")"):
            inner = lower[4:-1]
            parts = [p.strip() for p in inner.split(",")]
            if len(parts) == 3:
                return _normalize_rgb_tuple(parts)
        if raw.startswith("#"):
            raw = raw[1:]
        if len(raw) == 6:
            try:
                r = int(raw[0:2], 16)
                g = int(raw[2:4], 16)
                b = int(raw[4:6], 16)
                return _normalize_rgb_tuple((r, g, b))
            except Exception:
                return None
    return None


def _normalize_rgb_tuple(value):
    """
    统一 RGB 数据并校验范围。
    """
    try:
        r, g, b = value
        r, g, b = int(r), int(g), int(b)
        if not (0 <= r <= 255 and 0 <= g <= 255 and 0 <= b <= 255):
            return None
        return (r, g, b)
    except Exception:
        return None


def _apply_word_style(text_range, style):
    """
    将样式应用到 Word(.doc) 的 Range。
    """
    style = _normalize_style(style)
    if not style:
        return
    font_name = style.get("font_name")
    if font_name:
        text_range.Font.Name = str(font_name)

    font_size = _parse_font_size(style.get("font_size"))
    if font_size:
        text_range.Font.Size = float(font_size)

    color = _parse_color(style.get("color"))
    if color:
        text_range.Font.Color = _rgb_to_word_color(color)


def _rgb_to_word_color(rgb):
    """
    将 RGB 转换为 Word 颜色值。
    """
    r, g, b = rgb
    return int(r) + (int(g) << 8) + (int(b) << 16)


def _clear_document(doc):
    for paragraph in list(doc.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def _remove_paragraphs(doc, content):
    removed = 0
    for paragraph in list(doc.paragraphs):
        if content in paragraph.text:
            p = paragraph._element
            p.getparent().remove(p)
            removed += 1
    return removed


def _resolve_extension(path):
    lower = str(path).lower()
    if lower.endswith(".doc"):
        return ".doc"
    return ".docx"


def _create_doc_file_with_word(path, content, style=None):
    """
    使用 Word 创建 .doc 文件，并支持基础格式设置。
    """
    try:
        import win32com.client
    except Exception as e:
        return {"success": False, "reason": "doc_not_supported", "error": str(e), "message": "缺少 win32com，无法处理 .doc 文件"}
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Add()
        if content is not None:
            doc.Content.Text = _content_to_plain_text(content)
        _apply_word_style(doc.Content, style)
        doc.SaveAs(path, FileFormat=0)
        return {"success": True, "path": path, "message": "创建成功"}
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def _read_doc_with_word(path):
    try:
        import win32com.client
    except Exception as e:
        raise RuntimeError(f"缺少 win32com，无法处理 .doc 文件: {str(e)}")
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path, ReadOnly=True)
        text = doc.Content.Text or ""
        return text.replace("\r", "\n")
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def _update_doc_with_word(path, content, mode, style=None):
    """
    使用 Word 更新 .doc 文件内容，并支持基础格式设置。
    """
    try:
        import win32com.client
    except Exception as e:
        return {"success": False, "reason": "doc_not_supported", "error": str(e), "message": "缺少 win32com，无法处理 .doc 文件"}
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path, ReadOnly=False)
        text = doc.Content.Text or ""
        content_text = _content_to_plain_text(content) if content is not None else ""
        if mode == "replace":
            new_text = "" if content is None else content_text
        elif mode == "append":
            if content is None:
                return {"success": False, "reason": "content_empty", "message": "追加内容不能为空"}
            suffix = "" if text.endswith("\n") or not text else "\n"
            new_text = text + suffix + content_text
        elif mode == "remove":
            if not content:
                return {"success": False, "reason": "content_empty", "message": "要删除的内容不能为空"}
            if content_text not in text:
                return {"success": False, "reason": "content_not_found", "message": "未找到要删除的内容"}
            new_text = text.replace(content_text, "")
        else:
            return {"success": False, "reason": "mode_invalid", "message": "不支持的更新模式"}
        doc.Content.Text = new_text
        _apply_word_style(doc.Content, style)
        doc.Save()
        return {"success": True, "path": path, "message": "内容已更新", "mode": mode}
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
